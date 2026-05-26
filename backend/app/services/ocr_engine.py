import os
import cv2
import numpy as np
import pytesseract
from PIL import Image
from pypdf import PdfReader
import docx
import logging
from app.core.config import settings

logger = logging.getLogger(__name__)

# Configure Tesseract path for Windows if present
tesseract_env = os.getenv("TESSERACT_CMD")
if tesseract_env:
    pytesseract.pytesseract.tesseract_cmd = tesseract_env
elif os.name == 'nt':
    # Standard Windows installation paths
    default_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
    ]
    for path in default_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            logger.info(f"Tesseract found at: {path}")
            break

class OCREngine:
    @staticmethod
    def preprocess_image(image_path: str) -> str:
        """
        Preprocess image for optimal OCR extraction using OpenCV.
        1. Grayscale
        2. Noise reduction (Bilateral filter)
        3. Thresholding (Otsu's binarization)
        4. Deskewing (Skew correction)
        """
        try:
            # Read image
            img = cv2.imread(image_path)
            if img is None:
                raise ValueError("Could not read image from path")

            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

            # Noise reduction
            filtered = cv2.bilateralFilter(gray, 9, 75, 75)

            # Thresholding
            thresh = cv2.threshold(filtered, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

            # Deskewing (skew correction)
            coords = np.column_stack(np.where(thresh > 0))
            angle = 0.0
            if len(coords) > 0:
                angle = cv2.minAreaRect(coords)[-1]
                # Adjust angle representation
                if angle < -45:
                    angle = -(90 + angle)
                else:
                    angle = -angle
                    
            if abs(angle) > 0.5:
                (h, w) = img.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                thresh = cv2.warpAffine(thresh, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

            # Save preprocessed image back to same path (or temp path)
            processed_path = image_path.replace(".", "_processed.")
            cv2.imwrite(processed_path, thresh)
            return processed_path
        except Exception as e:
            logger.error(f"Image preprocessing failed: {e}. Using raw image.")
            return image_path

    @staticmethod
    def extract_text_from_image(image_path: str, preprocess: bool = True) -> str:
        """Run OCR on image file."""
        working_path = image_path
        if preprocess:
            working_path = OCREngine.preprocess_image(image_path)
            
        try:
            text = pytesseract.image_to_string(Image.open(working_path))
            
            # Cleanup processed file if it was created
            if preprocess and working_path != image_path and os.path.exists(working_path):
                os.remove(working_path)
                
            return text.strip()
        except Exception as e:
            logger.error(f"Tesseract OCR failed: {e}. Verify Tesseract binary is installed and configured.")
            # Fallback to simple OCR or prompt warning
            return f"[OCR Error: Tesseract not found or failed to execute. Details: {e}]"

    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> tuple[str, bool]:
        """
        Extract text from PDF.
        Returns:
            extracted_text (str): Combined text.
            is_scanned (bool): True if PDF was scanned (relied on OCR), False if selectable text.
        """
        try:
            reader = PdfReader(pdf_path)
            full_text = []
            
            # Try to extract text directly first
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    full_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            combined_text = "\n".join(full_text).strip()
            
            # If the extracted text is suspiciously small, treat it as a scanned PDF and run OCR
            if len(combined_text) < 100 * len(reader.pages):
                logger.info(f"PDF {os.path.basename(pdf_path)} appears to be scanned. Running OCR pipeline...")
                return OCREngine.ocr_pdf_pages(pdf_path), True
                
            return combined_text, False
        except Exception as e:
            logger.error(f"Selectable PDF read failed: {e}. Falling back to OCR.")
            return OCREngine.ocr_pdf_pages(pdf_path), True

    @staticmethod
    def ocr_pdf_pages(pdf_path: str) -> str:
        """Convert PDF pages to images and run OCR on each page."""
        try:
            from pdf2image import convert_from_path
            pages = convert_from_path(pdf_path, dpi=200)
            pdf_text = []
            
            for i, page in enumerate(pages):
                # Save page to temp image file
                temp_image_path = f"{pdf_path}_page_{i + 1}.png"
                page.save(temp_image_path, "PNG")
                
                # Preprocess and extract
                page_text = OCREngine.extract_text_from_image(temp_image_path, preprocess=True)
                pdf_text.append(f"--- Page {i + 1} ---\n{page_text}")
                
                # Clean up temp image
                if os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
                    
            return "\n".join(pdf_text)
        except Exception as e:
            logger.error(f"pdf2image processing failed: {e}. Ensure poppler-utils is installed.")
            # Fallback to direct raw PDF read or report error
            return f"[PDF-OCR Error: Ensure Poppler and Tesseract are installed. details: {e}]"

    @staticmethod
    def extract_text_from_docx(docx_path: str) -> str:
        """Extract text from MS Word document."""
        try:
            doc = docx.Document(docx_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    full_text.append(paragraph.text)
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return f"[DOCX Read Error: {e}]"

    @staticmethod
    def extract_text_from_txt(txt_path: str) -> str:
        """Read standard text files."""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT read failed: {e}")
            return f"[TXT Read Error: {e}]"
